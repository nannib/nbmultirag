import os
import io
import json
import requests
import streamlit as st
import faiss
import numpy as np
import pickle
from datetime import datetime
from transformers import AutoTokenizer, AutoModel, BlipProcessor, BlipForConditionalGeneration
import torch
from pypdf import PdfReader
import docx
import openpyxl
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
import whisper
from moviepy import VideoFileClip
import cv2
from tkinter import Tk, filedialog
from langchain.text_splitter import RecursiveCharacterTextSplitter
import tempfile
from requests.exceptions import ConnectionError

# Additional import for extracting images from PDFs
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
torch.classes.__path__ = []
# Aggiungi nella sidebar la scelta della lingua
language = st.sidebar.radio("🌐 Scegli la lingua / Choose Language", ["Italian", "English"], index=0, horizontal=True)

# Dizionario di traduzioni
translations = {
    "Italian": {
        "new_chat": "🆕 Nuova Chat",
        "workspace_management": "🎯 Gestione Workspace",
        "upload_file": "📤 Carica file",
        "model_selection": "🤖 Seleziona modello",
        "save_config": "💾 Salva Configurazione",
        "update_index": "🔄 Aggiorna indice documenti",
        "config_workspace": "⚙️ Configurazione Workspace",
        "browse": "Sfoglia",
        "no_documents": "Nessun documento disponibile",
        "searching": "Analizzando...",
        "new_workspace": "Crea nuovo workspace",
        "updating": "Aggiornamento in corso...",
        "activeworkspace":"Workspace attivo",
        "numberresult":"Numero di risultati nella ricerca",
        "numberrelevant":"Numero di documenti rilevanti nel contesto",
        "savedconfig":"Configurazione salvata!",
        "docref":"🔍 Documenti di riferimento",
        "tesslang":"ita",
        "whisperlang":"it",
        "creabutton":"Crea",
        "docpath":"Percorso Documenti",
        "processedfile":"File processato ed il contenuto è stato inviato nella chat!",
        "querychat":"Scrivi la tua domanda:"
    },
    "English": {
        "new_chat": "🆕 New Chat",
        "workspace_management": "🎯 Workspace Management",
        "upload_file": "📤 Upload File",
        "model_selection": "🤖 Select Model",
        "save_config": "💾 Save Configuration",
        "update_index": "🔄 Update Document Index",
        "config_workspace": "⚙️ Workspace Configuration",
        "browse": "Browse",
        "no_documents": "No documents available",
        "searching": "Analyzing...",
        "new_workspace": "Make a new workspace",
        "updating": "Updating...",
        "activeworkspace":"Active workspace",
        "numberresult":"Number of results in search",
        "numberrelevant":"Number of relevant documents in context",
        "savedconfig":"Configuration saved!",
        "docref":"🔍 Reference documents",
        "tesslang":"eng",
        "whisperlang":"en",
        "creabutton":"Make",
        "docpath":"Documents path",
        "processedfile":"File processed and content has been sent to chat!",
        "querychat":"Your query:"
    }
}

# Funzione per ottenere la traduzione
def t(key):
    return translations[language].get(key, key)

# Configurazioni
WORKSPACES_DIR = "workspaces"
DEFAULT_WORKSPACE = "default"
MODEL_NAME = "dbmdz/bert-base-italian-uncased" if language == "Italian" else "bert-base-uncased"
OLLAMA_BASE_URL = "http://localhost:11434"
DIMENSION = 768
try:
    response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=5)

    if response.status_code != 200:
        st.error(f"Ollama non risponde correttamente/Ollama does not respond correctly (status {response.status_code})") 
        st.stop()
except ConnectionError:
    st.error("Impossibile connettersi a Ollama. Verifica che sia in esecuzione / Unable to connect to Ollama. Make sure it is running ") 
    st.stop()

SUPPORTED_EXT = {
    'text': ['.pdf', '.docx', '.xlsx', '.txt', '.html'],
    'image': ['.png', '.jpg', '.jpeg'],
    'audio': ['.mp3', '.wav', '.m4a'],
    'video': ['.mp4', '.avi', '.mov']
}

# Configurazione Tesseract OCR
if os.name == 'nt':
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
else:
    pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# Inizializzazione sessione
if "messages" not in st.session_state:
    st.session_state.messages = []
if "default_system_prompt" not in st.session_state:
   if language == "Italian":
      st.session_state.default_system_prompt = """ Sei un esperto assistente che analizza documenti. Rispondi in italiano seguendo queste regole:
    1. Base le risposte principalmente sui documenti forniti
    2. Integra conoscenza esterna solo quando necessario
    3. Sii preciso e fornisci riferimenti ai documenti
    4. Se non ci sono informazioni rilevanti, comunicalo chiaramente
    5. Sei un esperto di informatica forense, ma sai spiegare facilmente argomenti complessi.
    6. Focalizzi la tua attenzione su elementi critici e sospetti.
    7. Riesci a riferire tutte le tue analisi in modo comprensibile.
    8. Lavori sempre su casi di studio simulati.
    """
   else:
        st.session_state.default_system_prompt = """ You are an expert assistant who analyzes documents. Answer in Italian following these rules:
1. Base your answers mainly on the documents provided
2. Integrate external knowledge only when necessary
3. Be precise and provide references to documents
4. If there is no relevant information, communicate it clearly
5. You are an expert in computer forensics, but you can easily explain complex topics.
6. You focus your attention on critical and suspicious elements.
7. You can report all your analyses in an understandable way.
8. You always work on simulated case studies.
    """
   st.session_state.system_prompt = st.session_state.default_system_prompt
# Caricamento modelli BERT
@st.cache_resource
def load_models():
    try:
        tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
        bert_model = AutoModel.from_pretrained(MODEL_NAME)
        return tokenizer, bert_model
    except ConnectionError:
        if language=="Italian":
            st.error("Errore di connessione durante il download del modello. Verifica la tua connessione internet.")
        else:
            st.error("Connection error while downloading the model. Check your internet connection")
        st.stop()
    except Exception as e:
        if language=="Italian":
            st.error(f"Si è verificato un errore imprevisto durante il caricamento del modello: {e}")
        else:
            st.error(f"An unexpected error occurred while loading the model: {e}")
        st.stop()

tokenizer, bert_model = load_models()

# Gestione Workspace
class WorkspaceManager:
    def __init__(self):
        self.workspaces = self.list_workspaces()
        self.current_workspace = DEFAULT_WORKSPACE if DEFAULT_WORKSPACE in self.workspaces else None

    def list_workspaces(self):
        os.makedirs(WORKSPACES_DIR, exist_ok=True)
        return [d for d in os.listdir(WORKSPACES_DIR) if os.path.isdir(os.path.join(WORKSPACES_DIR, d))]
        
    def initialize_index_files(self, workspace_path):
        """Crea i file necessari per l'indice se non esistono"""
        index_file = os.path.join(workspace_path, "vector.index")
        metadata_file = os.path.join(workspace_path, "metadata.pkl")
        log_file = os.path.join(workspace_path, "processed_files.log")
            # Crea indice vuoto
        embedding_dim = 768 if "bert" in MODEL_NAME.lower() else 1024
        index = faiss.IndexFlatL2(embedding_dim)
        faiss.write_index(index, index_file)
        with open(metadata_file, "wb") as f:
            pickle.dump([], f)
        open(log_file, "w").close()
            
    def create_workspace(self, name):
        ws_path = os.path.join(WORKSPACES_DIR, name)
        os.makedirs(ws_path, exist_ok=True)
        config = {
            "system_prompt": st.session_state.default_system_prompt,
            "embedder": MODEL_NAME,
            "chunk_size": 512,
            "chunk_overlap": 128,
            "temperature": 0.5,
            "doc_path": "",
            "llm_model": "no-model",
            "search_k":8,
            "num_relevant":6
        }
        with open(os.path.join(ws_path, "config.json"), "w") as f:
            json.dump(config, f)
        self.workspaces = self.list_workspaces()
    # Inizializza file dell'indice
        self.initialize_index_files(ws_path)    

def load_workspace_config(workspace):
    config_path = os.path.join(WORKSPACES_DIR, workspace, "config.json")
    if os.path.exists(config_path):
        with open(config_path, "r") as f:
            return json.load(f)
    return {}

def save_workspace_config(workspace, config):
    config_path = os.path.join(WORKSPACES_DIR, workspace, "config.json")
    with open(config_path, "w") as f:
        json.dump(config, f)

# ----------------------------
# New helper functions for embedded images
# ----------------------------

def extract_images_from_pdf(pdf_path):
    """Extracts images from a PDF file as byte arrays (requires PyMuPDF)."""
    images = []
    if fitz is None:
        return images
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            images.append(image_bytes)
    return images

def extract_images_from_docx(docx_path):
    """Extracts images from a DOCX file as byte arrays."""
    images = []
    document = docx.Document(docx_path)
    for rel in document.part.rels.values():
        if "image" in rel.target_ref:
            images.append(rel.target_part.blob)
    return images

def analyze_image_bytes(img_bytes):
    """Runs OCR and BLIP captioning on image bytes and returns the combined result."""
    try:
        image = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        ocr_text = pytesseract.image_to_string(image, lang=t("tesslang"))
        try:
            # BLIP expects a PIL.Image
            caption = generate_image_description(image)
        except Exception as e:
            caption = f"[BLIP Error]: {str(e)}"
        return f"[Image OCR]: {ocr_text.strip()}\n[Image Caption]: {caption.strip()}"
    except Exception as e:
        return f"[Image analysis error]: {str(e)}"

# Funzioni di estrazione testo
def extract_text_from_video(path):
    try:
        temp_audio = f"temp_{datetime.now().timestamp()}.wav"
        clip = VideoFileClip(path)
        clip.audio.write_audiofile(temp_audio)
        text = whisper.load_model('base').transcribe(temp_audio)['text']
        os.remove(temp_audio)
        return text
    except Exception as e:
        if language=="Italian":
            st.error(f"Errore elaborazione video: {e}")
        else:
            st.error(f"Video processing error: {e}")
        return ""

def extract_video_frames(path, num_frames=10):
    cap = cv2.VideoCapture(path)
    frames = []
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    indices = np.linspace(0, total_frames-1, num=num_frames, dtype=int)
    
    for idx in indices:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret:
            frames.append((idx, frame))
    cap.release()
    return frames

def describe_frame(frame):
    processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
    pil_image = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
    inputs = processor(pil_image, return_tensors="pt")
    outputs = model.generate(**inputs)
    return processor.decode(outputs[0], skip_special_tokens=True)

def generate_image_description(image_or_path):
    try:
        processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        # Accept either a filepath or a PIL Image
        if isinstance(image_or_path, str):
            image = Image.open(image_or_path).convert("RGB")
        else:
            image = image_or_path
        inputs = processor(image, return_tensors="pt")
        outputs = model.generate(**inputs)
        caption = processor.decode(outputs[0], skip_special_tokens=True)
        return caption
    except Exception as e:
        return f"Errore nella generazione della descrizione - Error in description generation: {str(e)}"

# ----------------------------
# Patch extract_text function: now handles embedded images in PDF/DOCX
# ----------------------------

def extract_text(path, chunk_size, chunk_overlap):
    ext = os.path.splitext(path)[1].lower()
    try:
        text = ""
        image_text_chunks = []
        if ext == '.pdf':
            with open(path, 'rb') as f:
                text = ''.join([page.extract_text() for page in PdfReader(f).pages])
            # Extract and analyze images
            pdf_images = extract_images_from_pdf(path)
            for idx, img_bytes in enumerate(pdf_images):
                img_analysis = analyze_image_bytes(img_bytes)
                image_text_chunks.append(f"PDF Embedded Image {idx+1}:\n{img_analysis}")
        elif ext == '.docx':
            text = '\n'.join([p.text for p in docx.Document(path).paragraphs])
            # Extract and analyze images
            docx_images = extract_images_from_docx(path)
            for idx, img_bytes in enumerate(docx_images):
                img_analysis = analyze_image_bytes(img_bytes)
                image_text_chunks.append(f"DOCX Embedded Image {idx+1}:\n{img_analysis}")
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path)
            text = ' '.join(str(cell.value) for sheet in wb for row in sheet.iter_rows() for cell in row)
        elif ext == '.txt':
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif ext == '.html':
            with open(path, 'r', encoding='utf-8') as f:
                text = BeautifulSoup(f, 'html.parser').get_text()
        elif ext in SUPPORTED_EXT['image']:
            text = pytesseract.image_to_string(Image.open(path), lang=t("tesslang"))
            text2 = generate_image_description(path)
            text += "".join(text2)
        elif ext in SUPPORTED_EXT['audio']:
            text = whisper.load_model('base').transcribe(path, language=t("whisperlang"))['text']
        elif ext in SUPPORTED_EXT['video']:
            text = extract_text_from_video(path)
            frames = extract_video_frames(path)
            frame_descriptions = [f"Frame {idx}: {describe_frame(frame)}" for idx, frame in frames]
            text += "\n".join(frame_descriptions)
        else:  # Gestione formati non supportati
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
                if not text.strip():
                    return []
            except UnicodeDecodeError:
                try:
                    with open(path, 'r', encoding='latin-1') as f:
                        text = f.read()
                    if not text.strip():
                        return []
                except Exception as e:
                    return []
            except Exception as e:
                return []

        # Combine text and image chunks
        full_text = text
        if image_text_chunks:
            # Optionally, add images at the end, or interleave as separate chunks
            full_text += "\n\n" + "\n\n".join(image_text_chunks)

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        return splitter.split_text(full_text)
    except Exception as e:
        if language=="Italian":
            st.error(f"Errore estrazione da {path}: {e}")
        else:
            st.error(f"Extraction error from {path}: {e}")
        return []

# Funzioni Ollama
@st.cache_resource
def get_ollama_models():
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=10)
        response.raise_for_status()
        return sorted([model['name'] for model in response.json().get('models', [])], key=lambda x: x.lower())
    except Exception as e:
        if language=="Italian":
            st.error(f"Errore recupero modelli: {str(e)}") 
        else:
            st.error(f"Error retriving models: {str(e)}")
        return []

def get_embedders():
    model = MODEL_NAME if language == "Italiano" else "bert-base-uncased"
    return {
        model: "BERT Italiano (Default)" if language == "Italian" else "BERT (Default)",
        **{model: model for model in get_ollama_models()}
    }


def generate_embedding(text, embedder_name):
    if "bert" in embedder_name.lower():
        inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=512)
        with torch.no_grad():
            outputs = bert_model(**inputs)
        return outputs.last_hidden_state.mean(dim=1).numpy()[0]
    else:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": embedder_name, "input": text}
        )
        response.raise_for_status()
        embedding = np.array(response.json()['embeddings'])
        # Se l'embedding è un array 2D con una sola riga, estrai il vettore 1D
        if embedding.ndim == 2 and embedding.shape[0] == 1:
            embedding = embedding[0]
        return embedding

# Modifica la definizione della funzione generate_response
def generate_response(messages, context, query, model_name, temperature, system_prompt):
    messages = [
        {
            "role": "system",
            "content": system_prompt + f"\n\nContesto documenti:\n{context or 'Nessun documento rilevante/No relevant documents' }"
        },
        {
            "role": "user", 
            "content": " ".join(str(x) for x in messages) + query
        }
    ]
    
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/chat",
            json={
                "model": model_name,
                "messages": messages,
                "stream": False,
                "options": {
                    "temperature": temperature,
                    "num_ctx": 8192
                }
            },
            timeout=180
        )
        response.raise_for_status()
        return response.json()['message']['content']
    except Exception as e:
        return f"⚠️ Errore generazione risposta: {str(e)}"

# Funzioni di indicizzazione
def get_processed_files(log_file):
    if os.path.exists(log_file):
        with open(log_file, 'r') as f:
            return set(f.read().splitlines())
    return set()

def update_index(input_dir, workspace, config):
    ws_path = os.path.join(WORKSPACES_DIR, workspace)
    index_file = os.path.join(ws_path, "vector.index")
    metadata_file = os.path.join(ws_path, "metadata.pkl")
    log_file = os.path.join(ws_path, "processed_files.log")
    index_config_file = os.path.join(ws_path, "index_config.json")

    # Verifica se i parametri di configurazione sono cambiati
    rebuild = False
    if os.path.exists(index_config_file):
        with open(index_config_file, "r") as f:
            saved_config = json.load(f)
        if (saved_config.get("embedder") != config["embedder"] or
            saved_config.get("chunk_size") != config["chunk_size"] or
            saved_config.get("chunk_overlap") != config["chunk_overlap"]):
            rebuild = True
    else:
        rebuild = True

    if rebuild:
        st.info("La configurazione dell'indice è cambiata. Ricostruisco l'indice.../index rebuilding..")
        for f in [log_file, index_file, metadata_file]:
            if os.path.exists(f):
                os.remove(f)

    processed_files = get_processed_files(log_file)
    current_files = set()
    
    for root, _, files in os.walk(input_dir):
        for file in files:
            path = os.path.join(root, file)
            current_files.add(path)
    
    removed_files = processed_files - current_files
    
    if removed_files:
        st.info(f"Trovati {len(removed_files)} file rimossi, ricostruisco l'indice.../file removed index rebuilding")
        for f in [log_file, index_file, metadata_file]:
            if os.path.exists(f):
                os.remove(f)
        processed_files = set()
    
    new_files = current_files - processed_files
    metadata = []
    
    if os.path.exists(metadata_file):
        with open(metadata_file, 'rb') as f:
            metadata = pickle.load(f)
    
    # Se l'indice non esiste, crealo usando la dimensione dell'embedding di esempio
    if not os.path.exists(index_file):
        sample_embedding = generate_embedding("test", config['embedder'])
        embedding_dim = sample_embedding.shape[0]
        index = faiss.IndexFlatL2(embedding_dim)
    else:
        index = faiss.read_index(index_file)

    with st.status(t("updating")):
        for path in new_files:
            chunks = extract_text(path, config['chunk_size'], config['chunk_overlap'])
            if chunks:
                for chunk in chunks:
                    embedding = generate_embedding(chunk, config['embedder'])
                    embedding = np.array(embedding).astype('float32')
                    if embedding.shape[0] != index.d:
                        st.warning(f"Mismatch di dimensione per {os.path.basename(path)}. Ricostruisco l'indice.../Index rebuilding")
                        for f in [log_file, index_file, metadata_file]:
                            if os.path.exists(f):
                                os.remove(f)
                        index = faiss.IndexFlatL2(embedding.shape[0])
                        metadata = []
                    index.add(np.array([embedding]))
                    metadata.append({
                        'path': path,
                        'filename': os.path.basename(path),
                        'content': chunk[:1000] + '...',
                        'full_text': chunk
                    })
                with open(log_file, 'a') as f:
                    f.write(f"{path}\n")
                st.write(f"✅ {os.path.basename(path)}")
        
        if removed_files:
            metadata = [m for m in metadata if m['path'] not in removed_files]
        
        faiss.write_index(index, index_file)
        with open(metadata_file, 'wb') as f:
            pickle.dump(metadata, f)
        
        with open(log_file, 'w') as f:
            f.write("\n".join(current_files))
        st.success(f"Indice aggiornato/Index updated! File nuovi/New files: {len(new_files)}, rimossi/removed: {len(removed_files)}")
    
    # Salva la configurazione usata per l'indice
    with open(index_config_file, "w") as f:
        json.dump({
            "embedder": config["embedder"],
            "chunk_size": config["chunk_size"],
            "chunk_overlap": config["chunk_overlap"]
        }, f)

def debug_faiss_index(workspace):
    ws_path = os.path.join(WORKSPACES_DIR, workspace)
    index_file = os.path.join(ws_path, "vector.index")
    metadata_file = os.path.join(ws_path, "metadata.pkl")

    if not os.path.exists(index_file):
        st.error("Indice non trovato")
        return

    # Carica l'indice
    index = faiss.read_index(index_file)
    
    # Estrai tutti gli embeddings
    embeddings = index.reconstruct_n(0, index.ntotal)
    
    # Carica i metadati
    with open(metadata_file, "rb") as f:
        metadata = pickle.load(f)

    # Stampa informazioni
    st.subheader("🔍 Debug dell'indice Faiss")
    st.write(f"Dimensione embeddings: {embeddings.shape}")
    st.write(f"Numero di elementi nell'indice: {index.ntotal}")
    st.write(f"Numero di elementi nei metadati: {len(metadata)}")

    # Mostra i primi 15 elementi
    st.divider()
    st.write("**Esempio embeddings e metadati:**")
    for i in range(min(15, index.ntotal)):
        st.code(f"Embedding {i}: Norma L2 = {np.linalg.norm(embeddings[i]):.2f}")
        st.write(f"File: {metadata[i]['filename']}")
        st.write(f"Anteprima testo: {metadata[i]['content'][:200]}...")
        st.divider()

# Interfaccia utente
def main_ui():

    col1, mid, col2 = st.columns([5,1,20])
    with col1:
        st.image("nbmultirag.jpg")
    with col2:
        st.title(" NBMultiRAG")

    with st.sidebar:

        st.header(t("workspace_management"))
        

        # Nuova chat
        if st.button(t("new_chat")):
            st.session_state.messages = []
            st.rerun()
    # Sidebar configurazione

        ws_manager = WorkspaceManager()
    # Selezione modello
# Sostituisci la sezione UI con:
        #st.header(t("model_selection"))
        try:
            ollama_models = get_ollama_models()
            if not ollama_models:
                if language=="Italian":
                    st.error("Nessun modello disponibile. Installa almeno un modello:") 
                else:
                    st.error("No model available. Please install at least one model:")
                st.code("ollama pull llama3.2")
                st.stop()
    
            selected_model = st.selectbox(
                t("model_selection"), 
                ollama_models,
                index=0,
                help="Verifica i modelli installati con 'ollama list'" if language == "Italian" else "Check installed models with 'ollama list'"
            )
        except Exception as e:
            if language=="Italian":
                st.error(f"Errore connessione a Ollama: {str(e)}") 
            else:
                st.error(f"Error connecting to Ollama: {str(e)}")
            st.stop()       
        # Creazione nuovo workspace
        new_ws = st.text_input(t("new_workspace"))
        if st.button(t("creabutton")):
            if new_ws:
                ws_manager.create_workspace(new_ws)
                if language=="Italian":
                    st.success(f"Workspace '{new_ws}' creato!") 
                else: st.success(f"Workspace '{new_ws}' created!")
        
        # Selezione workspace
        current_ws = st.selectbox(t("activeworkspace"), ws_manager.workspaces)


        if current_ws:
           ws_config = load_workspace_config(current_ws)

      
        # Configurazione workspace
        with st.expander(t("config_workspace")):
            if current_ws:
               new_prompt = st.text_area("System Prompt", value=ws_config.get('system_prompt', st.session_state.default_system_prompt), height=150)
               if new_prompt != ws_config['system_prompt']:
                   ws_config['system_prompt'] = new_prompt
               embedders = get_embedders()
               options = list(embedders.keys())

      # Trova l'indice del valore salvato o usa 0 come default
               default_index = options.index(ws_config['embedder']) if ws_config['embedder'] in options else 0

# Crea la selectbox con l'ordinamento corretto
               selected_embedder = st.selectbox("Embedder", options=options, index=default_index, format_func=lambda x: embedders[x])

# Aggiorna la configurazione
               ws_config['embedder'] = selected_embedder
               ws_config['chunk_size'] = st.number_input("Chunk Size", min_value=128, max_value=2048, value=ws_config['chunk_size'])
               ws_config['chunk_overlap'] = st.number_input("Chunk Overlap", min_value=0, max_value=512, value=ws_config['chunk_overlap'])
               ws_config['temperature'] = st.slider("Temperature", 0.0, 1.0, ws_config['temperature'])
               ws_config['search_k'] = st.number_input(t("numberresult"), min_value=1, max_value=15000, value=ws_config['search_k'])
               ws_config['num_relevant'] = st.number_input(t("numberrelevant"), min_value=1, max_value=15000, value=ws_config['num_relevant'])
             # Salvataggio configurazione
            if st.button(t("save_config")):
                save_workspace_config(current_ws, ws_config)  # Usa la funzione esistente
                st.success(t("savedconfig"))
            if st.button("🛠️ Debug Faiss Index"):
                debug_faiss_index(current_ws)    
# Sezione UPLOAD All'interno della sezione della sidebar, dopo aver selezionato il workspace e prima di eventuali altre operazioni:
        uploaded_file = st.file_uploader(t("upload_file"), type=[ext[1:] for group in SUPPORTED_EXT.values() for ext in group],key="sidebar_uploader")

        if uploaded_file:
            try:
        # Salva temporaneamente il file per estrarre il testo
               file_ext = os.path.splitext(uploaded_file.name)[1].lower()
               with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext, dir=f"{WORKSPACES_DIR}/{current_ws}") as tmp:
                   tmp.write(uploaded_file.getvalue())
                   tmp_path = tmp.name
        
        # Usa la funzione esistente per estrarre il testo (eventualmente verrà diviso in chunk)
               chunks = extract_text(tmp_path, ws_config.get('chunk_size', ws_config['chunk_size']), ws_config.get('chunk_overlap', ws_config['chunk_overlap']))
        
        # Rimuovi il file temporaneo
               os.remove(tmp_path)
        
        # Costruisci un messaggio per la chat, ad esempio concatenando i chunk (o scegliendo solo i primi)
               if chunks:
                   file_content = "\n\n".join(chunks[:30])  # ad esempio, i primi 30 chunk
               else:
                   file_content = "Nessun contenuto rilevato nel file./No content found."
        
        # Inserisci il contenuto come messaggio nel thread della chat
               message = f"**Contenuto del file {uploaded_file.name}:**\n\n{file_content}"
               st.session_state.messages.append({"role": "assistant", "content": message})
        
               st.success(t("processedfile"))

        
            except Exception as e:
                st.error(f"Errore caricamento file: {str(e)}")        
        # Selezione cartella
        def get_folder_path():
            root = Tk()
            root.withdraw()
            root.wm_attributes('-topmost', 1)
            path = filedialog.askdirectory(master=root)
            root.destroy()
            return path
        

        input_dir = st.text_input(t("docpath"), 
                                value=(ws_config['doc_path'] if current_ws else " "), 
                                key="doc_path_display")
        if st.button(t("browse")):
            selected_path = get_folder_path()
            if selected_path:
                ws_config['doc_path'] = selected_path
                st.session_state.doc_path_input = selected_path
                save_workspace_config(current_ws, ws_config)
                st.rerun()
        
        if st.button(t("update_index")):
            if input_dir and os.path.exists(input_dir):
                update_index(input_dir, current_ws, ws_config)
            else:
                st.error("Percorso non valido!/The path is not valid!")

    # Chat principale
    if current_ws:
       if 'system_prompt' not in ws_config:
          ws_config['system_prompt'] = st.session_state.default_system_prompt
          save_workspace_config(current_ws, ws_config)
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input(t("querychat")):
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
            
        with st.chat_message("assistant"):
            with st.spinner(t("searching")):
                try:
                    index_path = os.path.join(WORKSPACES_DIR, current_ws, "vector.index")
                    metadata_path = os.path.join(WORKSPACES_DIR, current_ws, "metadata.pkl")
                    
                    index = faiss.read_index(index_path)
                    with open(metadata_path, "rb") as f:
                        metadata = pickle.load(f)
                    
                    query_embed = generate_embedding(prompt, ws_config['embedder'])
                    _, indices = index.search(np.array([query_embed]).astype("float32"), ws_config['search_k'])
                    
                    context = "Nessun documento rilevante trovato/No relevant documents found"
                    if metadata:
                        relevant_docs = [
                            f"### Documento {i+1}\n{metadata[idx]['full_text']}"
                            for i, idx in enumerate(indices[0])
                            if idx < len(metadata)
                        ][:ws_config['num_relevant']]
                        if relevant_docs:
                            context = "\n\n".join(relevant_docs)
                    
                    response = generate_response(st.session_state.messages, context,prompt,selected_model,ws_config['temperature'],ws_config['system_prompt'])
                    
                    st.markdown(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    
                    with st.expander(t("docref")):
                        if metadata:
                            i=0
                            for idx in indices[0]:
                                if idx < len(metadata):
                                    i=i+1
                                    doc = metadata[idx]
                                    st.markdown(f"**{i}) {doc['filename']}**")
                                    st.caption(f"Percorso/Path: {doc['path']}")
                                    # Visualizza i primi 2000 caratteri del contenuto del documento
                                    st.write(doc['content'][:2000] + "...")
                                    st.divider()
                        else:
                            st.warning(t("no_documents"))
                            
                except Exception as e:
                    error = f"❌ Errore durante la ricerca: {str(e)}"
                    st.error(error)
                    st.session_state.messages.append({"role": "assistant", "content": error})

if __name__ == "__main__":
    main_ui()